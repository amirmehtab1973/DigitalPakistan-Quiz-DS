import streamlit as st
import pandas as pd
from datetime import datetime
import re
import io
import uuid
import json
import os
import time

# Import libraries with error handling
try:
    from PyPDF2 import PdfReader
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    st.warning("PyPDF2 not installed. PDF uploads will not work.")

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    st.warning("python-docx not installed. DOCX uploads will not work.")

# NLP Libraries with error handling
try:
    import nltk
    # Download required NLTK data
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt', quiet=True)
    try:
        nltk.data.find('corpora/stopwords')
    except LookupError:
        nltk.download('stopwords', quiet=True)
    try:
        nltk.data.find('taggers/averaged_perceptron_tagger')
    except LookupError:
        nltk.download('averaged_perceptron_tagger', quiet=True)
    
    from nltk.tokenize import sent_tokenize, word_tokenize
    from nltk.corpus import stopwords
    from nltk import pos_tag
    from nltk.chunk import ne_chunk
    from nltk.tree import Tree
    
    NLTK_SUPPORT = True
except ImportError:
    NLTK_SUPPORT = False
    st.warning("NLTK not installed. Advanced text processing will be limited.")

try:
    import spacy
    # Load spaCy model
    try:
        nlp = spacy.load("en_core_web_sm")
        SPACY_SUPPORT = True
    except OSError:
        st.warning("spaCy English model not found. Please run: python -m spacy download en_core_web_sm")
        SPACY_SUPPORT = False
except ImportError:
    SPACY_SUPPORT = False
    st.warning("spaCy not installed. Advanced NLP features will not be available.")

# File paths for persistence
QUIZZES_FILE = "quizzes.json"
STUDENT_RECORDS_FILE = "student_records.json"
COUNTER_FILE = "counter.json"

# Authentication credentials
ADMIN_CREDENTIALS = {
    "admin": "Admin123"
}

# Global variables to store data
quizzes_dict = {}
student_records = []
quiz_counter = 0

def load_data():
    """Load data from JSON files"""
    global quizzes_dict, student_records, quiz_counter
    
    try:
        # Load quizzes
        if os.path.exists(QUIZZES_FILE):
            with open(QUIZZES_FILE, 'r', encoding='utf-8') as f:
                quizzes_dict = json.load(f)
        else:
            quizzes_dict = {}
        
        # Load student records
        if os.path.exists(STUDENT_RECORDS_FILE):
            with open(STUDENT_RECORDS_FILE, 'r', encoding='utf-8') as f:
                student_records = json.load(f)
        else:
            student_records = []
        
        # Load counter
        if os.path.exists(COUNTER_FILE):
            with open(COUNTER_FILE, 'r', encoding='utf-8') as f:
                counter_data = json.load(f)
                quiz_counter = counter_data.get('quiz_counter', 0)
        else:
            quiz_counter = 0
            
    except Exception as e:
        st.error(f"Error loading data: {str(e)}")
        # Initialize empty data if loading fails
        quizzes_dict = {}
        student_records = []
        quiz_counter = 0

def save_quizzes():
    """Save quizzes to JSON file"""
    try:
        with open(QUIZZES_FILE, 'w', encoding='utf-8') as f:
            json.dump(quizzes_dict, f, indent=2, ensure_ascii=False)
    except Exception as e:
        st.error(f"Error saving quizzes: {str(e)}")

def save_student_records():
    """Save student records to JSON file"""
    try:
        with open(STUDENT_RECORDS_FILE, 'w', encoding='utf-8') as f:
            json.dump(student_records, f, indent=2, ensure_ascii=False)
    except Exception as e:
        st.error(f"Error saving student records: {str(e)}")

def save_counter():
    """Save counter to JSON file"""
    try:
        with open(COUNTER_FILE, 'w', encoding='utf-8') as f:
            json.dump({'quiz_counter': quiz_counter}, f, indent=2)
    except Exception as e:
        st.error(f"Error saving counter: {str(e)}")

def authenticate_user(username, password):
    """Authenticate admin/teacher user"""
    if username in ADMIN_CREDENTIALS and ADMIN_CREDENTIALS[username] == password:
        return True, "Authentication successful!"
    else:
        return False, "Invalid username or password!"

def extract_key_entities_nltk(text):
    """Extract key entities using NLTK"""
    if not NLTK_SUPPORT:
        return []
    
    try:
        # Tokenize and POS tag
        words = word_tokenize(text)
        pos_tags = pos_tag(words)
        
        # Extract nouns and proper nouns
        key_entities = []
        for word, pos in pos_tags:
            if pos in ['NN', 'NNS', 'NNP', 'NNPS'] and len(word) > 2:  # Nouns
                key_entities.append(word.lower())
        
        return list(set(key_entities))[:10]  # Return top 10 unique entities
    
    except Exception as e:
        st.error(f"NLTK entity extraction error: {str(e)}")
        return []

def extract_key_entities_spacy(text):
    """Extract key entities using spaCy"""
    if not SPACY_SUPPORT:
        return []
    
    try:
        doc = nlp(text)
        entities = []
        
        # Extract named entities
        for ent in doc.ents:
            if ent.label_ in ['PERSON', 'ORG', 'GPE', 'PRODUCT', 'EVENT', 'WORK_OF_ART']:
                entities.append(ent.text.lower())
        
        # Extract important nouns
        nouns = [token.lemma_.lower() for token in doc 
                if token.pos_ in ['NOUN', 'PROPN'] 
                and len(token.text) > 3
                and not token.is_stop]
        
        # Combine and get unique entities
        all_entities = list(set(entities + nouns))
        return all_entities[:15]  # Return top 15 unique entities
    
    except Exception as e:
        st.error(f"spaCy entity extraction error: {str(e)}")
        return []

def extract_key_concepts(text):
    """Extract key concepts using combined NLP approaches"""
    key_concepts = []
    
    # Use spaCy if available
    if SPACY_SUPPORT:
        spacy_entities = extract_key_entities_spacy(text)
        key_concepts.extend(spacy_entities)
    
    # Use NLTK as fallback or supplement
    if NLTK_SUPPORT or not key_concepts:
        nltk_entities = extract_key_entities_nltk(text)
        key_concepts.extend(nltk_entities)
    
    # Fallback to simple word frequency if NLP fails
    if not key_concepts:
        words = re.findall(r'\b[a-zA-Z]{4,}\b', text.lower())
        stop_words = set(stopwords.words('english')) if NLTK_SUPPORT else set()
        filtered_words = [word for word in words if word not in stop_words]
        from collections import Counter
        common_words = Counter(filtered_words).most_common(10)
        key_concepts = [word for word, count in common_words]
    
    return list(set(key_concepts))  # Return unique concepts

def generate_question_from_sentence(sentence, key_concepts, question_type):
    """Generate a specific type of question from a sentence"""
    
    if SPACY_SUPPORT:
        doc = nlp(sentence)
        # Extract subject and main verb
        subjects = [token.text for token in doc if token.dep_ in ["nsubj", "nsubjpass"]]
        verbs = [token.lemma_ for token in doc if token.pos_ == "VERB"]
        objects = [token.text for token in doc if token.dep_ in ["dobj", "attr", "pobj"]]
    else:
        # Simple fallback
        words = word_tokenize(sentence) if NLTK_SUPPORT else sentence.split()
        subjects = words[:3]  # Simple approximation
        verbs = []
        objects = []
    
    main_subject = subjects[0] if subjects else "the text"
    main_verb = verbs[0] if verbs else "discusses"
    main_object = objects[0] if objects else "concepts"
    
    question_templates = {
        'comprehension': [
            f"What is the main idea expressed in this sentence?",
            f"What does this sentence primarily discuss?",
            f"What is the key point being made here?"
        ],
        'detail': [
            f"According to the text, what is mentioned about {main_subject}?",
            f"What specific information is provided regarding {main_object}?",
            f"What details are given about {main_subject}?"
        ],
        'inference': [
            f"What can be inferred from this statement about {main_subject}?",
            f"Based on this information, what conclusion can be drawn about {main_object}?",
            f"What implication does this statement have for understanding {main_subject}?"
        ],
        'application': [
            f"How does this information about {main_subject} apply in practice?",
            f"What practical significance does this statement about {main_object} have?",
            f"How might this concept of {main_subject} be used?"
        ]
    }
    
    import random
    template_type = question_type if question_type in question_templates else random.choice(list(question_templates.keys()))
    template = random.choice(question_templates[template_type])
    
    return template

def generate_distractors(correct_answer, key_concepts, sentence_context):
    """Generate plausible distractors for MCQ options"""
    distractors = []
    
    # Type 1: Related but incorrect concepts
    if key_concepts:
        related = [concept for concept in key_concepts if concept != correct_answer]
        if related:
            distractors.append(f"Focuses on {random.choice(related)} instead")
    
    # Type 2: Opposite or contrasting ideas
    opposites = ["opposite perspective", "contrasting view", "different approach", "alternative interpretation"]
    distractors.append(f"Describes the {random.choice(opposites)}")
    
    # Type 3: Overgeneralization
    generalizations = ["broader context without specifics", "general principles only", "overview without details"]
    distractors.append(f"Provides a {random.choice(generalizations)}")
    
    # Type 4: Specific but wrong detail
    specifics = ["technical specifications", "historical background", "methodological details", "comparative analysis"]
    distractors.append(f"Emphasizes {random.choice(specifics)}")
    
    # Ensure we have exactly 3 distractors
    while len(distractors) < 3:
        distractors.append(f"Discusses related but different aspects")
    
    return distractors[:3]

def generate_mcqs_from_text(text, num_questions=10):
    """Generate high-quality MCQs from descriptive text using NLP"""
    try:
        # Preprocess text
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Extract key concepts using NLP
        key_concepts = extract_key_concepts(text)
        
        # Split into meaningful sentences
        if SPACY_SUPPORT:
            doc = nlp(text)
            sentences = [sent.text.strip() for sent in doc.sents if len(sent.text.strip()) > 40]
        elif NLTK_SUPPORT:
            sentences = sent_tokenize(text)
            sentences = [s.strip() for s in sentences if len(s.strip()) > 40]
        else:
            # Fallback to regex
            sentences = re.split(r'[.!?]+', text)
            sentences = [s.strip() for s in sentences if len(s.strip()) > 40]
        
        if not sentences:
            st.warning("No meaningful sentences found in the text.")
            return []
        
        # Limit number of questions
        num_questions = min(num_questions, len(sentences))
        questions = []
        
        # Question types for variety
        question_types = ['comprehension', 'detail', 'inference', 'application']
        
        import random
        
        for i in range(num_questions):
            sentence = sentences[i]
            question_type = question_types[i % len(question_types)]
            
            # Generate question
            question_text = generate_question_from_sentence(sentence, key_concepts, question_type)
            
            # Generate correct answer based on question type
            if question_type == 'comprehension':
                correct_answer = f"The main idea about the central concept"
            elif question_type == 'detail':
                correct_answer = f"Specific information provided in the text"
            elif question_type == 'inference':
                correct_answer = f"Logical conclusion based on the evidence"
            else:  # application
                correct_answer = f"Practical implications and applications"
            
            # Generate distractors
            distractors = generate_distractors(correct_answer, key_concepts, sentence)
            
            # Combine options
            all_options = [correct_answer] + distractors
            random.shuffle(all_options)
            
            # Find new position of correct answer
            correct_index = all_options.index(correct_answer)
            
            questions.append({
                'question_text': question_text,
                'options': all_options,
                'correct_answer': correct_index,
                'auto_generated': True,
                'source_sentence': sentence[:100] + "..." if len(sentence) > 100 else sentence
            })
        
        st.success(f"✅ Generated {len(questions)} MCQs using NLP analysis")
        if key_concepts:
            st.info(f"📊 Key concepts identified: {', '.join(key_concepts[:8])}")
        
        return questions
        
    except Exception as e:
        st.error(f"Error generating MCQs with NLP: {str(e)}")
        # Fallback to basic generation
        return generate_mcqs_from_text_basic(text, num_questions)

def generate_mcqs_from_text_basic(text, num_questions=5):
    """Fallback basic MCQ generation without NLP"""
    try:
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if len(s.strip()) > 30]
        
        if len(sentences) < num_questions:
            num_questions = len(sentences)
        
        questions = []
        import random
        
        for i in range(num_questions):
            if i >= len(sentences):
                break
                
            sentence = sentences[i]
            question_text = f"What is the primary focus of this statement: '{sentence[:80]}...'?"
            
            # Generate varied options
            options = [
                "The main concept and its implications",
                "Technical specifications and details", 
                "Historical context and background",
                "Comparative analysis with other topics"
            ]
            
            # Shuffle options
            random.shuffle(options)
            correct_index = random.randint(0, 3)  # Random correct answer
            
            questions.append({
                'question_text': question_text,
                'options': options,
                'correct_answer': correct_index,
                'auto_generated': True
            })
        
        return questions
        
    except Exception as e:
        st.error(f"Error in basic MCQ generation: {str(e)}")
        return []

# ... [The rest of your existing functions remain the same - parse_document, parse_mcqs_from_text, etc.]

def parse_document(file_obj, generate_mcqs=False):
    """Parse PDF or DOCX file and extract MCQs"""
    global quiz_counter
    
    try:
        if file_obj is None:
            return "Please select a file to upload."
        
        text = ""
        filename = file_obj.name
        
        try:
            if filename.endswith('.pdf'):
                if not PDF_SUPPORT:
                    return "PDF support not available. Please install PyPDF2."
                # Read PDF file
                pdf_reader = PdfReader(file_obj)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    
            elif filename.endswith('.docx'):
                if not DOCX_SUPPORT:
                    return "DOCX support not available. Please install python-docx."
                # Read DOCX file
                doc = Document(file_obj)
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"
            else:
                return "Unsupported file format. Please upload PDF or DOCX."
        
        except Exception as e:
            return f"Error reading file: {str(e)}"
        
        if not text.strip():
            return "No readable text found in the document."
        
        # Parse or generate MCQs
        if generate_mcqs:
            # Use NLP-enhanced generation
            questions = generate_mcqs_from_text(text, num_questions=8)  # Generate 8 questions by default
            if not questions:
                return "Could not generate MCQs from the document."
            message = f"✅ Successfully generated {len(questions)} MCQs using NLP analysis"
        else:
            questions = parse_mcqs_from_text(text)
            if not questions:
                return "No MCQs found in the document."
            message = f"✅ Successfully parsed {len(questions)} questions"
        
        # Create quiz entry
        quiz_counter += 1
        quiz_id = f"quiz_{quiz_counter}"
        quiz_title = os.path.basename(filename)
        
        quizzes_dict[quiz_id] = {
            'title': quiz_title,
            'questions': questions,
            'filename': filename,
            'enabled': False,
            'auto_generated': generate_mcqs,
            'duration_minutes': len(questions) * 2  # 2 minutes per question
        }
        
        # Save data
        save_quizzes()
        save_counter()
        
        return f"{message} - Duration: {len(questions) * 2} minutes"
    
    except Exception as e:
        return f"Unexpected error: {str(e)}"

# ... [The rest of your existing code remains unchanged]

# Import libraries with error handling
try:
    from PyPDF2 import PdfReader
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    st.warning("PyPDF2 not installed. PDF uploads will not work.")

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    st.warning("python-docx not installed. DOCX uploads will not work.")

# File paths for persistence
QUIZZES_FILE = "quizzes.json"
STUDENT_RECORDS_FILE = "student_records.json"
COUNTER_FILE = "counter.json"

# Authentication credentials
ADMIN_CREDENTIALS = {
    "admin": "Admin123"
}

# Global variables to store data
quizzes_dict = {}
student_records = []
quiz_counter = 0

def load_data():
    """Load data from JSON files"""
    global quizzes_dict, student_records, quiz_counter
    
    try:
        # Load quizzes
        if os.path.exists(QUIZZES_FILE):
            with open(QUIZZES_FILE, 'r', encoding='utf-8') as f:
                quizzes_dict = json.load(f)
        else:
            quizzes_dict = {}
        
        # Load student records
        if os.path.exists(STUDENT_RECORDS_FILE):
            with open(STUDENT_RECORDS_FILE, 'r', encoding='utf-8') as f:
                student_records = json.load(f)
        else:
            student_records = []
        
        # Load counter
        if os.path.exists(COUNTER_FILE):
            with open(COUNTER_FILE, 'r', encoding='utf-8') as f:
                counter_data = json.load(f)
                quiz_counter = counter_data.get('quiz_counter', 0)
        else:
            quiz_counter = 0
            
    except Exception as e:
        st.error(f"Error loading data: {str(e)}")
        # Initialize empty data if loading fails
        quizzes_dict = {}
        student_records = []
        quiz_counter = 0

def save_quizzes():
    """Save quizzes to JSON file"""
    try:
        with open(QUIZZES_FILE, 'w', encoding='utf-8') as f:
            json.dump(quizzes_dict, f, indent=2, ensure_ascii=False)
    except Exception as e:
        st.error(f"Error saving quizzes: {str(e)}")

def save_student_records():
    """Save student records to JSON file"""
    try:
        with open(STUDENT_RECORDS_FILE, 'w', encoding='utf-8') as f:
            json.dump(student_records, f, indent=2, ensure_ascii=False)
    except Exception as e:
        st.error(f"Error saving student records: {str(e)}")

def save_counter():
    """Save counter to JSON file"""
    try:
        with open(COUNTER_FILE, 'w', encoding='utf-8') as f:
            json.dump({'quiz_counter': quiz_counter}, f, indent=2)
    except Exception as e:
        st.error(f"Error saving counter: {str(e)}")

def authenticate_user(username, password):
    """Authenticate admin/teacher user"""
    if username in ADMIN_CREDENTIALS and ADMIN_CREDENTIALS[username] == password:
        return True, "Authentication successful!"
    else:
        return False, "Invalid username or password!"

def generate_mcqs_from_text(text, num_questions=5):
    """Generate MCQs from descriptive text"""
    try:
        # Split text into sentences
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
        
        if len(sentences) < num_questions:
            num_questions = len(sentences)
        
        questions = []
        
        for i in range(num_questions):
            if i >= len(sentences):
                break
                
            sentence = sentences[i]
            words = sentence.split()
            
            if len(words) < 5:
                continue
                
            # Create question
            question_text = f"What is the main idea of: '{sentence[:100]}...'?"
            
            # Generate options
            options = [
                "The text discusses general concepts",
                f"The passage focuses on {words[2] if len(words) > 2 else 'key'} aspects",
                "It describes historical events",
                "The content explains technical details"
            ]
            
            questions.append({
                'question_text': question_text,
                'options': options,
                'correct_answer': 0,  # Default to first option
                'auto_generated': True
            })
        
        return questions
        
    except Exception as e:
        st.error(f"Error generating MCQs: {str(e)}")
        return []

def parse_document(file_obj, generate_mcqs=False):
    """Parse PDF or DOCX file and extract MCQs"""
    global quiz_counter
    
    try:
        if file_obj is None:
            return "Please select a file to upload."
        
        text = ""
        filename = file_obj.name
        
        try:
            if filename.endswith('.pdf'):
                if not PDF_SUPPORT:
                    return "PDF support not available. Please install PyPDF2."
                # Read PDF file
                pdf_reader = PdfReader(file_obj)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    
            elif filename.endswith('.docx'):
                if not DOCX_SUPPORT:
                    return "DOCX support not available. Please install python-docx."
                # Read DOCX file
                doc = Document(file_obj)
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"
            else:
                return "Unsupported file format. Please upload PDF or DOCX."
        
        except Exception as e:
            return f"Error reading file: {str(e)}"
        
        if not text.strip():
            return "No readable text found in the document."
        
        # Parse or generate MCQs
        if generate_mcqs:
            questions = generate_mcqs_from_text(text)
            if not questions:
                return "Could not generate MCQs from the document."
            message = f"✅ Successfully generated {len(questions)} MCQs"
        else:
            questions = parse_mcqs_from_text(text)
            if not questions:
                return "No MCQs found in the document."
            message = f"✅ Successfully parsed {len(questions)} questions"
        
        # Create quiz entry
        quiz_counter += 1
        quiz_id = f"quiz_{quiz_counter}"
        quiz_title = os.path.basename(filename)
        
        quizzes_dict[quiz_id] = {
            'title': quiz_title,
            'questions': questions,
            'filename': filename,
            'enabled': False,
            'auto_generated': generate_mcqs,
            'duration_minutes': len(questions)
        }
        
        # Save data
        save_quizzes()
        save_counter()
        
        return f"{message} - Duration: {len(questions)} minutes"
    
    except Exception as e:
        return f"Unexpected error: {str(e)}"

def parse_mcqs_from_text(text):
    """Extract MCQs from text"""
    questions = []
    
    try:
        # Clean the text
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\n+', '\n', text)
        
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # Check if this is a question line
            if '?' in line and len(line) > 10:
                question_text = line
                options = []
                i += 1
                
                # Look for options in subsequent lines
                option_count = 0
                while i < len(lines) and option_count < 4:
                    current_line = lines[i]
                    
                    # Check if this is an option (starts with A., B., etc.)
                    if re.match(r'^[A-D][\.\)]', current_line, re.IGNORECASE):
                        option_text = re.sub(r'^[A-D][\.\)]\s*', '', current_line, flags=re.IGNORECASE)
                        options.append(option_text.strip())
                        option_count += 1
                    # Stop if we find another question
                    elif '?' in current_line and len(current_line) > 10:
                        break
                    
                    i += 1
                
                if question_text and len(options) >= 2:
                    # Pad options to 4 if needed
                    while len(options) < 4:
                        options.append("")
                    
                    questions.append({
                        'question_text': question_text,
                        'options': options,
                        'correct_answer': None,
                        'auto_generated': False
                    })
                else:
                    i += 1
            else:
                i += 1
        
        return questions
    
    except Exception as e:
        st.error(f"Error parsing MCQs: {str(e)}")
        return []

def get_student_quiz_choices():
    """Get choices for student quiz dropdown"""
    choices = []
    for quiz_id, quiz_data in quizzes_dict.items():
        if quiz_data['enabled']:
            total_questions = len(quiz_data['questions'])
            correct_set = sum(1 for q in quiz_data['questions'] if q['correct_answer'] is not None)
            if correct_set == total_questions:
                duration = quiz_data.get('duration_minutes', total_questions)
                choices.append((quiz_id, f"📝 {quiz_data['title']} ({total_questions} questions, {duration} minutes)"))
    
    # If no choices, add a helpful message
    if not choices:
        choices.append(("", "❌ No quizzes available - teacher must enable quizzes and set answers"))
    
    return choices

def toggle_quiz_enabled(quiz_id):
    """Toggle quiz enabled/disabled status"""
    if quiz_id and quiz_id in quizzes_dict:
        quizzes_dict[quiz_id]['enabled'] = not quizzes_dict[quiz_id]['enabled']
        status = "enabled" if quizzes_dict[quiz_id]['enabled'] else "disabled"
        save_quizzes()
        return f"✅ Quiz {status} successfully!"
    return "❌ No quiz selected"

def submit_student_quiz(quiz_id, student_name, student_email, answers):
    """Process student quiz submission"""
    if not quiz_id or quiz_id not in quizzes_dict:
        return "❌ Quiz not found. Please refresh and try again."
    
    if not student_name.strip() or not student_email.strip():
        return "❌ Please enter your name and email."
    
    quiz = quizzes_dict[quiz_id]
    questions = quiz['questions']
    
    # Calculate score
    score = 0
    total = len(questions)
    
    for i, question in enumerate(questions):
        if i < len(answers) and answers[i] is not None and answers[i] == question['correct_answer']:
            score += 1
    
    # Store student record
    record = {
        'id': str(uuid.uuid4())[:8],
        'quiz_id': quiz_id,
        'quiz_title': quiz['title'],
        'student_name': student_name.strip(),
        'student_email': student_email.strip(),
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'score': score,
        'total_questions': total,
        'percentage': round((score / total) * 100, 2) if total > 0 else 0
    }
    
    student_records.append(record)
    save_student_records()
    
    # Return both the result HTML and the record for display
    result_html = f"""
    <div style="padding: 20px; border: 2px solid #4CAF50; border-radius: 10px; background: #f9fff9;">
        <h3 style="color: #4CAF50; text-align: center;">🎉 Quiz Completed!</h3>
        <div style="text-align: center; margin: 20px 0;">
            <h2>Score: {score}/{total} ({record['percentage']}%)</h2>
        </div>
        <div style="background: white; padding: 15px; border-radius: 8px;">
            <p><strong>Name:</strong> {student_name}</p>
            <p><strong>Email:</strong> {student_email}</p>
            <p><strong>Quiz:</strong> {quiz['title']}</p>
            <p><strong>Date/Time:</strong> {record['timestamp']}</p>
        </div>
    </div>
    """
    
    return result_html, record

def generate_student_report():
    """Generate Excel report of all student results"""
    if not student_records:
        return "❌ No student records found.", None
    
    # Create DataFrame
    data = []
    for record in student_records:
        data.append({
            'Student ID': record['id'],
            'Student Name': record['student_name'],
            'Student Email': record['student_email'],
            'Quiz Title': record['quiz_title'],
            'Date/Time': record['timestamp'],
            'Score': record['score'],
            'Total Questions': record['total_questions'],
            'Percentage': f"{record['percentage']}%"
        })
    
    df = pd.DataFrame(data)
    
    # Create Excel file in memory
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Student Results', index=False)
    
    output.seek(0)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"student_results_{timestamp}.xlsx"
    
    return output, filename

# Initialize session state
if 'authenticated' not in st.session_state:
    st.session_state.authenticated = False
if 'quiz_active' not in st.session_state:
    st.session_state.quiz_active = False
if 'current_quiz_id' not in st.session_state:
    st.session_state.current_quiz_id = None
if 'student_answers' not in st.session_state:
    st.session_state.student_answers = {}
if 'current_student_name' not in st.session_state:
    st.session_state.current_student_name = ""
if 'current_student_email' not in st.session_state:
    st.session_state.current_student_email = ""
if 'quiz_start_time' not in st.session_state:
    st.session_state.quiz_start_time = None
if 'quiz_duration' not in st.session_state:
    st.session_state.quiz_duration = None
if 'time_expired' not in st.session_state:
    st.session_state.time_expired = False
if 'auto_submitted' not in st.session_state:
    st.session_state.auto_submitted = False
if 'quiz_result' not in st.session_state:
    st.session_state.quiz_result = None
if 'last_refresh' not in st.session_state:
    st.session_state.last_refresh = 0
if 'force_refresh' not in st.session_state:
    st.session_state.force_refresh = False
if 'last_auto_refresh' not in st.session_state:
    st.session_state.last_auto_refresh = 0
if 'refresh_requested' not in st.session_state:
    st.session_state.refresh_requested = False

# Load data
load_data()

# Main application
st.set_page_config(
    page_title="Digital Pakistan Quiz Management System",
    page_icon="🎯",
    layout="wide"
)

# Custom CSS with improved timer and button positioning
st.markdown("""
<style>
    .main-header {
        text-align: center;
        color: #1f77b4;
        padding: 20px;
    }
    .timer-wrapper {
        position: fixed;
        top: 20px;
        right: 20px;
        z-index: 9999;
        text-align: center;
        background: white;
        padding: 15px;
        border-radius: 10px;
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
        border: 2px solid #e0e0e0;
        min-width: 200px;
    }
    .timer-container {
        background: #4CAF50;
        color: white;
        padding: 15px;
        border-radius: 10px;
        text-align: center;
        border: 2px solid #45a049;
        font-family: Arial, sans-serif;
        margin-bottom: 10px;
    }
    .timer-warning {
        background: #FF9800 !important;
        border-color: #ffb74d !important;
    }
    .timer-danger {
        background: #ff4444 !important;
        border-color: #ff6b6b !important;
    }
    .quiz-container {
        border: 2px solid #2196F3;
        border-radius: 10px;
        padding: 20px;
        margin: 10px 0;
        background: #f0f8ff;
    }
    .refresh-warning {
        background: #fff3cd;
        border: 1px solid #ffeaa7;
        border-radius: 5px;
        padding: 10px;
        margin: 10px 0;
        text-align: center;
    }
    .time-expired-warning {
        background: #f8d7da;
        border: 1px solid #f5c6cb;
        border-radius: 5px;
        padding: 15px;
        margin: 10px 0;
        text-align: center;
        color: #721c24;
    }
</style>
""", unsafe_allow_html=True)

# Header
st.markdown("""
<div class="main-header">
    <h1>🎯 Digital Pakistan Quiz Management System</h1>
    <p><strong>Teacher Panel:</strong> Upload quiz documents and set correct answers<br>
    <strong>Student Panel:</strong> Take quizzes and view results</p>
</div>
""", unsafe_allow_html=True)

# Improved auto-refresh and auto-submit logic
if st.session_state.quiz_active and st.session_state.quiz_start_time and st.session_state.quiz_duration:
    current_time = time.time()
    elapsed_time = current_time - st.session_state.quiz_start_time
    remaining_time = max(0, st.session_state.quiz_duration - elapsed_time)
    
    # Auto-refresh every 1 second to check for time expiration
    if current_time - st.session_state.last_refresh >= 1:
        st.session_state.last_refresh = current_time
        
        # Check if time expired and auto-submit
        if remaining_time <= 0 and not st.session_state.auto_submitted:
            st.session_state.time_expired = True
            st.session_state.auto_submitted = True
            st.session_state.force_refresh = True

# Handle manual refresh request
if st.session_state.refresh_requested:
    st.session_state.refresh_requested = False
    st.session_state.force_refresh = True

# Auto-refresh logic - check if 10 seconds have passed
current_time = time.time()
if (st.session_state.quiz_active and 
    st.session_state.last_auto_refresh and 
    current_time - st.session_state.last_auto_refresh >= 10):
    st.session_state.force_refresh = True
    st.session_state.last_auto_refresh = current_time

# Handle auto-submission
if st.session_state.auto_submitted and st.session_state.quiz_active:
    # Collect all answers
    answers = []
    if st.session_state.current_quiz_id in quizzes_dict:
        questions = quizzes_dict[st.session_state.current_quiz_id]['questions']
        for i in range(len(questions)):
            answer_key = f"q_{i}"
            answers.append(st.session_state.student_answers.get(answer_key))
        
        # Submit quiz
        result = submit_student_quiz(
            st.session_state.current_quiz_id,
            st.session_state.current_student_name,
            st.session_state.current_student_email,
            answers
        )
        
        if isinstance(result, tuple) and len(result) == 2:
            result_html, record = result
            # Store result in session state
            st.session_state.quiz_result = result_html
        
        # Reset quiz state but keep result
        st.session_state.quiz_active = False
        st.session_state.current_quiz_id = None
        st.session_state.student_answers = {}
        st.session_state.quiz_start_time = None
        st.session_state.quiz_duration = None
        st.session_state.time_expired = False
        st.session_state.auto_submitted = False
        st.session_state.last_auto_refresh = 0
        st.session_state.force_refresh = True

# Force refresh if needed
if st.session_state.get('force_refresh', False):
    st.session_state.force_refresh = False
    st.rerun()

# Create tabs
tab1, tab2 = st.tabs(["🎓 Student Panel", "👨‍🏫 Teacher Admin Panel"])

# Student Panel
with tab1:
    st.header("🎓 Student Panel")
    
    # Show quiz result if available
    if st.session_state.quiz_result:
        st.markdown(st.session_state.quiz_result, unsafe_allow_html=True)
        if st.button("Take Another Quiz"):
            st.session_state.quiz_result = None
            st.rerun()
    
    else:
        if st.button("🔄 Refresh Quiz List"):
            st.rerun()
        
        student_choices = get_student_quiz_choices()
        
        if not student_choices or student_choices[0][0] == "":
            st.warning("❌ No quizzes available. The teacher must enable quizzes and set correct answers first.")
        else:
            col1, col2 = st.columns(2)
            with col1:
                student_name = st.text_input("Your Name", placeholder="Enter your full name", key="student_name")
            with col2:
                student_email = st.text_input("Your Email", placeholder="Enter your email address", key="student_email")
            
            selected_quiz_option = st.selectbox(
                "Select Quiz to Take",
                options=[choice[0] for choice in student_choices],
                format_func=lambda x: dict(student_choices)[x],
                key="student_quiz_select"
            )
            
            if st.button("Start Quiz", type="primary", key="start_quiz_btn"):
                if not student_name.strip() or not student_email.strip():
                    st.error("❌ Please enter your name and email.")
                else:
                    st.session_state.quiz_active = True
                    st.session_state.current_quiz_id = selected_quiz_option
                    st.session_state.current_student_name = student_name
                    st.session_state.current_student_email = student_email
                    st.session_state.student_answers = {}
                    st.session_state.quiz_start_time = time.time()
                    if selected_quiz_option in quizzes_dict:
                        quiz = quizzes_dict[selected_quiz_option]
                        st.session_state.quiz_duration = quiz.get('duration_minutes', len(quiz['questions'])) * 60
                    st.session_state.time_expired = False
                    st.session_state.auto_submitted = False
                    st.session_state.quiz_result = None
                    st.session_state.last_refresh = time.time()
                    st.session_state.last_auto_refresh = time.time()
                    st.rerun()
            
            # Display active quiz
            if st.session_state.quiz_active and st.session_state.current_quiz_id == selected_quiz_option:
                quiz = quizzes_dict[st.session_state.current_quiz_id]
                questions = quiz['questions']
                duration_minutes = quiz.get('duration_minutes', len(questions))
                
                # Calculate remaining time
                if st.session_state.quiz_start_time and st.session_state.quiz_duration:
                    current_time = time.time()
                    elapsed_time = current_time - st.session_state.quiz_start_time
                    remaining_time = max(0, st.session_state.quiz_duration - elapsed_time)
                    minutes = int(remaining_time // 60)
                    seconds = int(remaining_time % 60)
                    
                    # Determine timer color
                    timer_class = ""
                    if remaining_time < 60:
                        timer_class = "timer-danger"
                    elif remaining_time < st.session_state.quiz_duration // 2:
                        timer_class = "timer-warning"
                    
                    # Calculate time until next auto-refresh
                    time_since_refresh = current_time - st.session_state.last_auto_refresh
                    time_until_refresh = max(0, 10 - time_since_refresh)
                    
                    # Display timer
                    timer_html = f"""
                    <div class="timer-wrapper">
                        <div class="timer-container {timer_class}">
                            <div style="font-size: 14px; font-weight: bold; margin-bottom: 5px;">⏰ QUIZ TIMER</div>
                            <div style="font-size: 18px; font-weight: bold; font-family: 'Courier New', monospace; margin-bottom: 3px;">
                                {minutes:02d}:{seconds:02d}
                            </div>
                            <div style="font-size: 11px; opacity: 0.9;">
                                {duration_minutes} minute quiz
                            </div>
                            <div style="font-size: 10px; margin-top: 8px; color: #666;">
                                Auto-refresh in: {int(time_until_refresh)}s
                            </div>
                        </div>
                    </div>
                    """
                    st.markdown(timer_html, unsafe_allow_html=True)
                
                # Add a working Streamlit refresh button in the main content area
                col1, col2, col3 = st.columns([2, 1, 2])
                with col2:
                    if st.button("🔄 Refresh Timer Now", key="manual_refresh_btn", type="primary", use_container_width=True):
                        st.session_state.refresh_requested = True
                        st.rerun()
                
                # Time expired warning
                if st.session_state.time_expired:
                    st.error("⏰ TIME'S UP! Your quiz is being submitted...")
                
                st.markdown(f"""
                <div class="quiz-container">
                    <h3>📝 Taking Quiz: {quiz['title']}</h3>
                    <p><strong>Total Questions:</strong> {len(questions)} | <strong>Time Allowed:</strong> {duration_minutes} minutes</p>
                    <p><em>Page auto-refreshes every 10 seconds. Use 'Refresh Timer Now' button above for immediate update.</em></p>
                </div>
                """, unsafe_allow_html=True)
                
                # Display questions
                for i, question in enumerate(questions):
                    st.subheader(f"Question {i+1}: {question['question_text']}")
                    
                    # Create options
                    options = []
                    for j, option in enumerate(question['options']):
                        if option.strip():
                            options.append(f"{chr(65+j)}: {option}")
                    
                    # Use radio buttons for answers
                    answer_key = f"q_{i}"
                    if answer_key not in st.session_state.student_answers:
                        st.session_state.student_answers[answer_key] = None
                    
                    selected_option = st.radio(
                        f"Select your answer for Question {i+1}:",
                        options=range(len(options)) if options else [],
                        format_func=lambda x: options[x] if x < len(options) else "Invalid",
                        key=answer_key,
                        index=st.session_state.student_answers[answer_key] if st.session_state.student_answers[answer_key] is not None else 0
                    )
                    
                    st.session_state.student_answers[answer_key] = selected_option
                    st.divider()
                
                # Submit button at the bottom
                st.markdown("---")
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    if st.button("Submit Quiz", type="primary", key="submit_quiz_btn", use_container_width=True):
                        # Collect all answers
                        answers = []
                        for i in range(len(questions)):
                            answer_key = f"q_{i}"
                            answers.append(st.session_state.student_answers.get(answer_key))
                        
                        # Submit quiz
                        result = submit_student_quiz(
                            st.session_state.current_quiz_id,
                            st.session_state.current_student_name,
                            st.session_state.current_student_email,
                            answers
                        )
                        
                        if isinstance(result, tuple) and len(result) == 2:
                            result_html, record = result
                            # Store result
                            st.session_state.quiz_result = result_html
                        
                        # Reset quiz state
                        st.session_state.quiz_active = False
                        st.session_state.current_quiz_id = None
                        st.session_state.student_answers = {}
                        st.session_state.quiz_start_time = None
                        st.session_state.quiz_duration = None
                        st.session_state.time_expired = False
                        st.session_state.auto_submitted = False
                        st.session_state.last_auto_refresh = 0
                        
                        st.rerun()

# Teacher Panel (same as before)
with tab2:
    st.header("👨‍🏫 Teacher Admin Panel")
    
    # Login section
    if not st.session_state.authenticated:
        st.subheader("🔐 Teacher/Admin Login")
        col1, col2 = st.columns(2)
        with col1:
            username = st.text_input("Username", placeholder="Enter username", key="teacher_username")
        with col2:
            password = st.text_input("Password", type="password", placeholder="Enter password", key="teacher_password")
        
        if st.button("Login", type="primary", key="login_btn"):
            success, message = authenticate_user(username, password)
            if success:
                st.session_state.authenticated = True
                st.success(message)
                st.rerun()
            else:
                st.error(message)
    else:
        st.success("✅ Logged in as Admin")
        
        # Upload section
        st.subheader("📤 Upload Quiz Document")
        uploaded_file = st.file_uploader("Upload PDF or DOCX file with MCQs", type=["pdf", "docx"], key="file_uploader")
        generate_mcqs = st.checkbox("🤖 Generate MCQs from descriptive text (for articles/descriptive documents)", key="generate_mcqs_cb")
        
        if st.button("Upload and Parse Quiz", type="primary", key="upload_btn"):
            if uploaded_file is not None:
                result = parse_document(uploaded_file, generate_mcqs)
                if result.startswith("✅"):
                    st.success(result)
                else:
                    st.error(result)
            else:
                st.error("Please select a file to upload.")
        
        # Edit quizzes section
        st.subheader("✏️ Set Correct Answers")
        if quizzes_dict:
            quiz_options = []
            for quiz_id, quiz_data in quizzes_dict.items():
                total_questions = len(quiz_data['questions'])
                correct_set = sum(1 for q in quiz_data['questions'] if q['correct_answer'] is not None)
                status = "✅" if correct_set == total_questions else "⚠️"
                enabled_status = "🟢" if quiz_data['enabled'] else "🔴"
                auto_gen = "🤖" if quiz_data.get('auto_generated', False) else "📝"
                duration = quiz_data.get('duration_minutes', total_questions)
                quiz_options.append((quiz_id, f"{enabled_status} {auto_gen} {status} {quiz_data['title']} ({correct_set}/{total_questions}) - {duration}min"))
            
            selected_quiz_id = st.selectbox(
                "Select Quiz to Edit",
                options=[option[0] for option in quiz_options],
                format_func=lambda x: dict(quiz_options)[x],
                key="edit_quiz_select"
            )
            
            if selected_quiz_id:
                quiz = quizzes_dict[selected_quiz_id]
                questions = quiz['questions']
                
                st.markdown(f"""
                <div style="padding: 15px; border: 2px solid #4CAF50; border-radius: 10px; background: #f9fff9; margin-bottom: 20px;">
                    <h4>🎯 Setting Correct Answers: {quiz['title']}</h4>
                    <p><strong>Total Questions:</strong> {len(questions)} | <strong>Duration:</strong> {quiz.get('duration_minutes', len(questions))} minutes</p>
                </div>
                """, unsafe_allow_html=True)
                
                # Display questions for setting correct answers
                saved_answers = []
                for i, question in enumerate(questions):
                    st.subheader(f"Question {i+1}: {question['question_text']}")
                    if question.get('auto_generated', False):
                        st.info("🤖 This question was auto-generated")
                    
                    # Create options
                    options = []
                    for j, option in enumerate(question['options']):
                        if option.strip():
                            options.append(f"{chr(65+j)}: {option}")
                    
                    # Current correct answer
                    current_answer = question['correct_answer']
                    
                    # Radio button for correct answer
                    correct_answer = st.radio(
                        f"Select correct answer for Question {i+1}",
                        options=range(len(options)) if options else [],
                        format_func=lambda x: options[x] if x < len(options) else "Invalid",
                        index=current_answer if current_answer is not None and current_answer < len(options) else 0,
                        key=f"teacher_{selected_quiz_id}_{i}"
                    )
                    
                    saved_answers.append(correct_answer)
                    st.divider()
                
                # Save all answers button
                if st.button("💾 Save All Answers", type="primary", key="save_answers_btn"):
                    for i, question in enumerate(questions):
                        if i < len(saved_answers):
                            question['correct_answer'] = saved_answers[i]
                    
                    save_quizzes()
                    st.success(f"✅ Saved all {len(questions)} answers successfully!")
                    st.rerun()
        
        # Enable/disable quizzes
        st.subheader("🎯 Enable/Disable Quizzes")
        if quizzes_dict:
            enable_quiz_options = []
            for quiz_id, quiz_data in quizzes_dict.items():
                total_questions = len(quiz_data['questions'])
                correct_set = sum(1 for q in quiz_data['questions'] if q['correct_answer'] is not None)
                status = "✅ Ready" if correct_set == total_questions else f"⚠️ {correct_set}/{total_questions} answers set"
                enable_quiz_options.append((quiz_id, f"{quiz_data['title']} - {status}"))
            
            selected_enable_quiz = st.selectbox(
                "Select Quiz to Toggle",
                options=[option[0] for option in enable_quiz_options],
                format_func=lambda x: dict(enable_quiz_options)[x],
                key="enable_quiz_select"
            )
            
            if selected_enable_quiz:
                current_status = quizzes_dict[selected_enable_quiz]['enabled']
                status_text = "🟢 ENABLED" if current_status else "🔴 DISABLED"
                st.write(f"Current status: **{status_text}**")
                
                if st.button("🔄 Toggle Quiz Status", key="toggle_btn"):
                    result = toggle_quiz_enabled(selected_enable_quiz)
                    if result.startswith("✅"):
                        st.success(result)
                        st.rerun()
                    else:
                        st.error(result)
        
        # Student records section
        st.subheader("📊 Student Results")
        
        if student_records:
            # Display student records using st.dataframe only
            st.write(f"**Total Records:** {len(student_records)}")
            
            # Create a clean dataframe for display
            display_data = []
            for record in sorted(student_records, key=lambda x: x['timestamp'], reverse=True)[:20]:
                display_data.append({
                    'Student Name': record['student_name'],
                    'Email': record['student_email'],
                    'Quiz': record['quiz_title'],
                    'Score': f"{record['score']}/{record['total_questions']} ({record['percentage']}%)",
                    'Date/Time': record['timestamp']
                })
            
            df_display = pd.DataFrame(display_data)
            st.dataframe(df_display, use_container_width=True)
            
            st.write(f"*Showing latest 20 of {len(student_records)} records*")
            
            # Download button
            if st.button("📥 Download Excel Report", key="download_btn"):
                excel_data, filename = generate_student_report()
                if excel_data and filename:
                    st.download_button(
                        label="⬇️ Click to Download Student Results (Excel)",
                        data=excel_data,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key="download_excel_btn"
                    )
                else:
                    st.error("Could not generate Excel report.")
        else:
            st.info("📝 No student records yet. Students need to take quizzes first.")
        
        # Logout button
        if st.button("🚪 Logout", key="logout_btn"):
            st.session_state.authenticated = False
            st.rerun()

# Footer
st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #666;">
    <p>🎯 <strong>Digital Pakistan Quiz Management System</strong> - Streamlining education through technology</p>
</div>
""", unsafe_allow_html=True)

